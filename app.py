import os
import subprocess
from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import JWTManager
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

app = Flask(__name__)
CORS(app)

# -------------------------------
# Environment variables
# -------------------------------
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'fallbacksecret')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///default.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
jwt = JWTManager(app)

# -------------------------------
# Health check
# -------------------------------
@app.route('/ping')
def ping():
    return jsonify({"status": "ok"})

# -------------------------------
# Root route
# -------------------------------
@app.route('/')
def home():
    return jsonify({"message": "Resume Screener Backend is running!"})

# -------------------------------
# Utility function
# -------------------------------
def calculate_similarity_with_keywords(resume_text, job_text):
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform([resume_text, job_text])
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

    resume_words = set(resume_text.lower().split())
    job_words = set(job_text.lower().split())
    common_words = list(resume_words.intersection(job_words))

    return round(similarity * 100, 2), common_words

# -------------------------------
# Upload resume (frontend expects /upload_resume)
# -------------------------------
@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    if 'resume' not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files['resume']
    text = ""

    try:
        filename = file.filename.lower()

        if filename.endswith(".pdf"):
            file.seek(0)
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        pil_image = page.to_image(resolution=300).original
                        ocr_text = pytesseract.image_to_string(pil_image)
                        if ocr_text.strip():
                            text += ocr_text + "\n"

        elif filename.endswith(".docx"):
            file.seek(0)
            doc = Document(file)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

        else:
            file.seek(0)
            text = file.read().decode('utf-8', errors='ignore')

    except Exception as e:
        return jsonify({"error": f"Failed to extract text: {str(e)}"}), 500

    return jsonify({"resume_text": text})

# -------------------------------
# Match routes
# -------------------------------
@app.route('/match', methods=['POST'])
def match_resume():
    data = request.json
    resume = data.get("resume", "")
    job = data.get("job", "")
    if not resume or not job:
        return jsonify({"error": "Resume or job description missing"}), 400
    score, keywords = calculate_similarity_with_keywords(resume, job)
    return jsonify({"score": score, "keywords": keywords})

@app.route('/match_multiple', methods=['POST'])
def match_multiple():
    data = request.json
    resume = data.get("resume", "")
    jobs = data.get("jobs", [])
    if not resume or not jobs:
        return jsonify({"error": "Resume or job descriptions missing"}), 400

    results = []
    for job in jobs:
        score, keywords = calculate_similarity_with_keywords(resume, job["description"])
        results.append({
            "title": job["title"],
            "score": score,
            "keywords": keywords
        })

    return jsonify({"results": results})

# -------------------------------
# Resume summary
# -------------------------------
@app.route('/resume_summary', methods=['POST'])
def resume_summary():
    data = request.json
    resume = data.get("resume", "")
    if not resume:
        return jsonify({"error": "Resume text missing"}), 400

    skills_list = [
        "python", "sql", "java", "c++", "tableau", "power bi", "excel",
        "kubernetes", "docker", "aws", "azure", "gcp",
        "prometheus", "grafana", "victoriametrics",
        "machine learning", "data science", "analytics",
        "digital marketing", "seo", "sem"
    ]

    resume_lower = resume.lower()
    matched_skills = [skill for skill in skills_list if skill in resume_lower]

    if matched_skills:
        summary = f"This resume highlights skills in: {', '.join(matched_skills)}"
    else:
        summary = "No specific technical skills detected in the resume."

    return jsonify({"summary": summary})

# -------------------------------
# Check Tesseract
# -------------------------------
@app.route('/check_tesseract')
def check_tesseract():
    try:
        version = subprocess.check_output(['tesseract', '--version']).decode('utf-8').split('\n')[0]
        return jsonify({"tesseract_version": version})
    except Exception as e:
        return jsonify({"error": str(e)})

# -------------------------------
# Run Flask app (local only)
# -------------------------------
if __name__ == '__main__':
    app.run(debug=True)
